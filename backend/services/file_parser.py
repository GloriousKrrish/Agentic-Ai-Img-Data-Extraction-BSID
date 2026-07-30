import os
import json
import csv
import io
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from PIL import Image, ImageOps
import pypdf
import docx

def parse_file_content(file_bytes: bytes, file_name: str, mime_type: str = "") -> dict:
    """
    Ingests any input file format and extracts text, metadata, images, or tabular structures.
    Supports: Images, PDF, DOCX, XLSX, CSV, JSON, XML, TXT, ZIP.
    """
    ext = Path(file_name).suffix.lower()
    
    # 1. Images
    if ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'] or 'image' in mime_type:
        try:
            image = Image.open(io.BytesIO(file_bytes))
            image = ImageOps.exif_transpose(image) # Auto-rotate based on EXIF
            width, height = image.size
            format_name = image.format or "JPEG"
            return {
                "file_type": "image",
                "format": format_name,
                "dimensions": f"{width}x{height}",
                "text_content": "",
                "raw_bytes": file_bytes,
                "has_vision": True
            }
        except Exception as e:
            return {"file_type": "image", "error": str(e), "raw_bytes": file_bytes, "has_vision": True}

    # 2. PDF Documents
    elif ext == '.pdf' or 'pdf' in mime_type:
        extracted_text = ""
        page_count = 0
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    extracted_text += f"\n--- Page {i+1} ---\n" + txt
        except Exception as e:
            extracted_text = f"PDF Read Error: {str(e)}"
            
        return {
            "file_type": "pdf",
            "page_count": page_count,
            "text_content": extracted_text.strip(),
            "raw_bytes": file_bytes,
            "has_vision": True # Can also be passed to Gemini Vision
        }

    # 3. Word Documents (.docx)
    elif ext == '.docx':
        extracted_text = ""
        tables_data = []
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            for p in doc.paragraphs:
                if p.text.strip():
                    extracted_text += p.text + "\n"
            for t in doc.tables:
                table_rows = []
                for row in t.rows:
                    table_rows.append([cell.text.strip() for cell in row.cells])
                tables_data.append(table_rows)
        except Exception as e:
            extracted_text = f"Docx Read Error: {str(e)}"

        return {
            "file_type": "docx",
            "text_content": extracted_text.strip(),
            "tables": tables_data,
            "has_vision": False
        }

    # 3b. Excel Workbooks (.xlsx, .xls)
    elif ext in ['.xlsx', '.xls'] or 'spreadsheet' in mime_type or 'excel' in mime_type:
        text_lines = []
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            sheet = wb.active
            for r in sheet.iter_rows(values_only=True):
                vals = [str(v).strip() for v in r if v is not None and str(v).strip()]
                if vals:
                    text_lines.append(" | ".join(vals))
            wb.close()
        except Exception as e:
            text_lines = [f"Excel Read Error: {str(e)}"]

        return {
            "file_type": "xlsx",
            "text_content": "\n".join(text_lines[:100]),
            "has_vision": False
        }

    # 4. CSV Files
    elif ext == '.csv' or 'csv' in mime_type:
        rows = []
        try:
            decoded = file_bytes.decode('utf-8', errors='ignore')
            reader = csv.reader(io.StringIO(decoded))
            for row in reader:
                rows.append(row)
        except Exception as e:
            rows = [["CSV Error", str(e)]]

        return {
            "file_type": "csv",
            "rows": rows,
            "text_content": "\n".join([", ".join(r) for r in rows[:50]]),
            "has_vision": False
        }

    # 5. JSON Files
    elif ext == '.json' or 'json' in mime_type:
        parsed_json = None
        try:
            decoded = file_bytes.decode('utf-8', errors='ignore')
            parsed_json = json.loads(decoded)
        except Exception as e:
            parsed_json = {"error": str(e)}

        return {
            "file_type": "json",
            "parsed_json": parsed_json,
            "text_content": json.dumps(parsed_json, indent=2)[:4000],
            "has_vision": False
        }

    # 6. XML Files
    elif ext == '.xml' or 'xml' in mime_type:
        text_content = ""
        try:
            decoded = file_bytes.decode('utf-8', errors='ignore')
            root = ET.fromstring(decoded)
            text_content = ET.tostring(root, encoding='utf-8', method='text').decode('utf-8')
        except Exception as e:
            text_content = f"XML Error: {str(e)}"

        return {
            "file_type": "xml",
            "text_content": text_content[:4000],
            "has_vision": False
        }

    # 7. ZIP Archives
    elif ext == '.zip':
        file_list = []
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                file_list = z.namelist()
        except Exception as e:
            file_list = [f"Zip Error: {str(e)}"]

        return {
            "file_type": "zip",
            "contained_files": file_list,
            "text_content": "ZIP Archive containing: " + ", ".join(file_list[:20]),
            "has_vision": False
        }

    # 8. Plain Text / Code / Fallback
    else:
        text_content = ""
        try:
            text_content = file_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            text_content = f"Binary file or unknown format: {str(e)}"

        return {
            "file_type": "text",
            "text_content": text_content[:4000],
            "has_vision": False
        }
